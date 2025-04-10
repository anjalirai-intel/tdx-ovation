import os
import docx
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from langchain.chains import RetrievalQA
from langchain_core.prompts import PromptTemplate
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyPDFLoader, DirectoryLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
import requests
from langchain_openai import AzureOpenAIEmbeddings, AzureChatOpenAI
from langchain_core.chat_history import InMemoryChatMessageHistory
from langchain_core.runnables.history import RunnableWithMessageHistory

load_dotenv()

os.environ["https_proxy"] = "http://proxy-dmz.intel.com:912"
os.environ["http_proxy"] = "http://proxy-dmz.intel.com:911"
os.environ["no_proxy"] = "localhost,127.0.0.1,::1,.openai.azure.com"

VECTORDB_PATH = "local_db"
DATA_FOLDER = "data"
HF_TOKEN=os.environ.get("HF_TOKEN")
chat_history = InMemoryChatMessageHistory()

PROMPT_TEMPLATE = """
    You are a highly knowledgeable and helpful assistant specializing in TDX (Trust Domain Extensions). 
    Your goal is to provide clear, concise, and accurate answers to user queries related to TDX. 

    Guidelines:
    1. Always base your answers on the provided context. If the context does not contain enough information, 
    respond with "I am not aware of this. Please send an email to the respective person for further support."
    2. Avoid making assumptions or providing incorrect information.
    3. Keep your responses professional and to the point.

    Context: {context}
    Question: {question}
"""

def read_pdf_files():
    loader = DirectoryLoader(DATA_FOLDER, glob="*.pdf", loader_cls=PyPDFLoader)
    documents = loader.load()
    return documents

def read_word_file(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def read_txt_file(file_path):
    with open(file_path, 'r') as file:
        text = file.read()
    return text

def read_html_file(file_path='', html_url=''):
    if file_path:
        with open(file_path, "r", encoding="utf-8") as file:
            html_content = file.read()
    elif html_url:
        response = requests.get(html_url, proxies=None, timeout=10)
        html_content = response.text
    soup = BeautifulSoup(html_content, "html.parser")
    text = soup.get_text()

    cleaned_text = ' '.join(text.split())
    return cleaned_text

def read_files():
    documents = []

    for root, _, files in os.walk(DATA_FOLDER):
        for file in files:
            file_path = os.path.join(root, file)
            file_extension = os.path.splitext(file)[1].lower()

            # Handle PDF files
            if file_extension == ".pdf":
                loader = PyPDFLoader(file_path)
                documents.extend(loader.load())

            # Handle Word files
            elif file_extension == ".docx":
                documents.append(Document(page_content=read_word_file(file_path), metadata={"source": file_path}))

            # Handle HTML files
            elif file_extension == ".html":
                documents.append(Document(page_content=read_html_file(file_path), metadata={"source": file_path}))

            # Handle plain text files
            elif file_extension == ".txt":
                documents.append(Document(page_content=read_txt_file(file_path), metadata={"source": file_path}))

    # Read html url from csv file
    with open(os.path.join(DATA_FOLDER, "html_links.csv"), encoding="utf-8-sig") as csv_file:
        for url in csv_file:
            if "github.com" in url:
                url = url.replace("github.com", "raw.githubusercontent.com")
            if "/blob/" in url:
                url = url.replace("/blob/", "/")
            documents.append(Document(page_content=read_html_file(html_url=url.strip()), metadata={"source": url}))
    
    return documents

def create_chunks(extracted_data):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    text_chunks = text_splitter.split_documents(extracted_data)
    return text_chunks

def create_vector_store(text_chunks):
    # embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    embeddings = AzureOpenAIEmbeddings(
        model=os.environ["OPENAI_EMBEDDING_MODEL"],
        api_key=os.environ["AZURE_OPENAI_KEY"],
        azure_endpoint=os.environ["AZURE_OPENAI_URL"],
        api_version=os.environ["OPENAI_EMBEDDING_VERSION"],
        chunk_size=1000
    )
    vector_db = FAISS.from_documents(text_chunks, embeddings)
    vector_db.save_local(VECTORDB_PATH)

def custom_prompt(template):
    prompt = PromptTemplate(template=template, input_variables=["context", "question"])
    return prompt

def load_llm():
    # llm = HuggingFaceEndpoint(
    #     repo_id="mistralai/Mistral-7B-Instruct-v0.3",
    #     task="text-generation",
    #     temperature=0.5,
    #     model_kwargs={
    #         "token": HF_TOKEN
    #     }
    # )
    llm = AzureChatOpenAI(
        model=os.environ["OPENAI_CHAT_MODEL"],
        api_version=os.environ["OPENAI_CHAT_VERSION"],
        api_key=os.environ["AZURE_OPENAI_KEY"],
        azure_endpoint=os.environ["AZURE_OPENAI_URL"],
        temperature=0.2,
    )
    return llm

def get_memory(_):
    return chat_history

def query_chatbot(prompt, vector_store):

    qa_chain = RetrievalQA.from_chain_type(
        llm=load_llm(),
        chain_type="stuff",
        retriever=vector_store.as_retriever(search_type="similarity", search_kwargs={'k': 3}),
        return_source_documents=False,
        chain_type_kwargs={'prompt': custom_prompt(PROMPT_TEMPLATE)}
    )

    qa_chain_with_history = RunnableWithMessageHistory(
        qa_chain,
        get_memory,
        input_messages_key="query",
        history_messages_key="chat_history"
    )

    response = qa_chain_with_history.invoke(
        {'query': prompt},
        config={"configurable": {"session_id": "rag-memory"}})

    return response['result']